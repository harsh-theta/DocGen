"""
Tests for export formatting consistency between PDF and DOCX formats.

This module tests the heading hierarchy and layout synchronization between
PDF and DOCX exports to ensure consistent formatting across both formats.
"""

import pytest
import os
import io
import tempfile
from pathlib import Path
from unittest.mock import patch, Mock, AsyncMock
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

from backend.services.export_manager import ExportFormatManager
from backend.services.document_formatter import DocumentFormatter

# Mock WeasyPrint if it's not available
try:
    from backend.services.weasyprint_generator import WeasyPrintGenerator, WEASYPRINT_AVAILABLE
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    WeasyPrintGenerator = Mock()

# Skip tests that require WeasyPrint
weasyprint_required = pytest.mark.skipif(
    not WEASYPRINT_AVAILABLE,
    reason="WeasyPrint or its dependencies are not available"
)


@pytest.fixture
def export_config():
    """Standard export configuration for consistent testing."""
    return {
        'font_family': 'Inter',
        'font_size': 12,
        'margins': {'top': 25, 'right': 25, 'bottom': 25, 'left': 25},
        'cover_page_mode': 'minimal',
        'table_style': {
            'border_color': '#dddddd',
            'border_width': '1px',
            'header_bg_color': '#f3f3f3',
            'cell_padding': '8px',
            'text_align': 'left'
        }
    }


@pytest.fixture
def sample_html_with_headings():
    """Sample HTML content with heading hierarchy for testing."""
    return """
    <h1>Main Document Title</h1>
    <p>This is the introduction paragraph.</p>
    <h2>Section 1</h2>
    <p>This is the first section content.</p>
    <h3>Subsection 1.1</h3>
    <p>This is subsection 1.1 content.</p>
    <h3>Subsection 1.2</h3>
    <p>This is subsection 1.2 content.</p>
    <h4>Subsubsection 1.2.1</h4>
    <p>This is subsubsection 1.2.1 content.</p>
    <h2>Section 2</h2>
    <p>This is the second section content.</p>
    <h3>Subsection 2.1</h3>
    <p>This is subsection 2.1 content.</p>
    <h2>Section 3</h2>
    <p>This is the third section content.</p>
    """


class TestExportFormatting:
    """Test cases for export formatting consistency between PDF and DOCX formats."""
    
    def test_heading_hierarchy_styling(self, export_config, sample_html_with_headings):
        """Test that heading hierarchy styling is consistent."""
        manager = ExportFormatManager(export_config)
        
        # Apply heading hierarchy styling
        styled_html = manager.apply_heading_hierarchy(sample_html_with_headings)
        
        # Parse the styled HTML
        soup = BeautifulSoup(styled_html, 'html.parser')
        
        # Verify each heading level has appropriate styling
        for level in range(1, 5):  # h1 to h4
            tag = f'h{level}'
            headings = soup.find_all(tag)
            
            # Verify we have headings at this level
            assert len(headings) > 0, f"No {tag} headings found"
            
            for heading in headings:
                # Verify style attribute exists
                assert 'style' in heading.attrs, f"{tag} is missing style attribute"
                
                # Verify font family
                assert f"font-family:'{export_config['font_family']}'" in heading['style']
                
                # Verify font size based on heading level
                expected_size = manager.heading_sizes[tag]
                assert f"font-size:{expected_size}pt" in heading['style']
                
                # Verify margins
                assert "margin-top" in heading['style']
                assert "margin-bottom" in heading['style']
    
    def test_docx_style_dict_generation(self, export_config):
        """Test that DOCX style dictionary is generated correctly."""
        manager = ExportFormatManager(export_config)
        
        # Get DOCX style dictionary
        docx_style = manager.get_docx_style_dict()
        
        # Verify font family
        assert docx_style['font_family'] == export_config['font_family']
        
        # Verify font size
        assert docx_style['font_size'] == export_config['font_size']
        
        # Verify margins (converted to DOCX units)
        for key in ['top', 'right', 'bottom', 'left']:
            # PDF uses mm, DOCX uses twips (1/20th of a point)
            # The conversion factor is applied in the ExportFormatManager
            assert docx_style['margins'][key] == pytest.approx(export_config['margins'][key] * 2.835, rel=0.1)
        
        # Verify heading sizes
        for level in range(1, 7):
            tag = f'h{level}'
            assert tag in docx_style['heading_sizes']
            assert docx_style['heading_sizes'][tag] == manager.heading_sizes[tag]
        
        # Verify table style
        assert 'table_style' in docx_style
        assert docx_style['table_style']['border_color'] == export_config['table_style']['border_color']
        assert docx_style['table_style']['header_bg_color'] == export_config['table_style']['header_bg_color']
    
    def test_css_template_generation(self, export_config):
        """Test that CSS template is generated correctly."""
        manager = ExportFormatManager(export_config)
        
        # Get CSS template
        css = manager.get_css_template()
        
        # Verify font family
        assert f"font-family: '{export_config['font_family']}'" in css
        
        # Verify font size
        assert f"font-size: {export_config['font_size']}pt" in css
        
        # Verify margins
        margin_values = f"{export_config['margins']['top']}mm {export_config['margins']['right']}mm {export_config['margins']['bottom']}mm {export_config['margins']['left']}mm"
        assert f"margin: {margin_values}" in css or f"margin:{margin_values}" in css
        
        # Verify heading styles
        for level in range(1, 7):
            tag = f'h{level}'
            assert tag + " {" in css or tag + "{" in css
    
    def test_prepare_html_for_export(self, export_config, sample_html_with_headings):
        """Test that HTML is properly prepared for export."""
        manager = ExportFormatManager(export_config)
        
        # Prepare HTML for export
        prepared_html = manager.prepare_html_for_export(sample_html_with_headings)
        
        # Verify structure
        assert "<!DOCTYPE html>" in prepared_html
        assert "<html>" in prepared_html
        assert "<head>" in prepared_html
        assert "<style>" in prepared_html
        assert "</style>" in prepared_html
        assert "</head>" in prepared_html
        assert "<body>" in prepared_html
        assert "</body>" in prepared_html
        assert "</html>" in prepared_html
        
        # Parse the prepared HTML
        soup = BeautifulSoup(prepared_html, 'html.parser')
        
        # Verify headings have styles
        for level in range(1, 5):  # h1 to h4
            tag = f'h{level}'
            headings = soup.find_all(tag)
            for heading in headings:
                assert 'style' in heading.attrs
        
        # Verify tables have styles
        tables = soup.find_all('table')
        for table in tables:
            assert 'style' in table.attrs
    
    def test_docx_heading_structure(self, export_config, sample_html_with_headings):
        """Test that headings are properly structured in DOCX."""
        # Create a DOCX document
        doc = DocxDocument()
        
        # Parse the HTML
        soup = BeautifulSoup(sample_html_with_headings, 'html.parser')
        
        # Process elements (simplified version of export_docx logic)
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
            if element.name.startswith('h'):
                # Extract heading level
                level = int(element.name[1])
                # Add heading with appropriate level
                doc.add_heading(element.get_text(strip=True), level=level)
            elif element.name == 'p':
                # Add paragraph
                doc.add_paragraph(element.get_text(strip=True))
        
        # Save to buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Read back and verify
        doc = DocxDocument(docx_buffer)
        
        # Count headings by level
        h1_count = 0
        h2_count = 0
        h3_count = 0
        h4_count = 0
        
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1':
                h1_count += 1
            elif paragraph.style.name == 'Heading 2':
                h2_count += 1
            elif paragraph.style.name == 'Heading 3':
                h3_count += 1
            elif paragraph.style.name == 'Heading 4':
                h4_count += 1
        
        # Verify heading counts match the original HTML
        assert h1_count == 1  # Main Document Title
        assert h2_count == 3  # Section 1, 2, 3
        assert h3_count == 3  # Subsection 1.1, 1.2, 2.1
        assert h4_count == 1  # Subsubsection 1.2.1
    
    def test_get_export_config(self, export_config):
        """Test that export configuration is correctly generated for different formats."""
        manager = ExportFormatManager(export_config)
        
        # Get PDF export config
        pdf_config = manager.get_export_config('pdf')
        
        # Verify PDF config
        assert 'css' in pdf_config
        assert 'font_paths' in pdf_config
        assert 'margins' in pdf_config
        assert 'font_family' in pdf_config
        assert 'font_size' in pdf_config
        
        # Get DOCX export config
        docx_config = manager.get_export_config('docx')
        
        # Verify DOCX config
        assert 'font_family' in docx_config
        assert 'font_size' in docx_config
        assert 'margins' in docx_config
        assert 'heading_sizes' in docx_config
        assert 'table_style' in docx_config
        
        # Test invalid format
        with pytest.raises(ValueError, match="Unsupported export format"):
            manager.get_export_config('invalid')


if __name__ == "__main__":
    # Run tests with pytest
    pytest.main([__file__, "-v"])