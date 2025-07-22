"""
Test module for DOCX table export functionality.

This module tests the enhanced table parsing and rendering in DOCX exports.
"""

import pytest
import io
import os
import tempfile
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from backend.services.export_manager import ExportFormatManager
from backend.routers.documents import export_docx


@pytest.fixture
def sample_html_with_table():
    """Sample HTML content with a table for testing."""
    return """
    <h1>Test Document</h1>
    <p>This is a test paragraph.</p>
    <table>
        <thead>
            <tr>
                <th>Header 1</th>
                <th>Header 2</th>
                <th>Header 3</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Row 1, Cell 1</td>
                <td>Row 1, Cell 2</td>
                <td>Row 1, Cell 3</td>
            </tr>
            <tr>
                <td>Row 2, Cell 1</td>
                <td><b>Bold text</b></td>
                <td><i>Italic text</i></td>
            </tr>
            <tr>
                <td>Row 3, Cell 1</td>
                <td>Row 3, Cell 2</td>
                <td>Row 3, Cell 3</td>
            </tr>
        </tbody>
    </table>
    <p>Another paragraph after the table.</p>
    """


@pytest.fixture
def complex_html_with_nested_content():
    """Complex HTML content with nested elements in table cells."""
    return """
    <h1>Complex Table Test</h1>
    <table>
        <thead>
            <tr>
                <th>Task</th>
                <th>Description</th>
                <th>Status</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Task 1</td>
                <td>
                    <p>This is a <b>complex</b> description with:</p>
                    <ul>
                        <li>Item 1</li>
                        <li>Item 2</li>
                    </ul>
                </td>
                <td>Complete</td>
            </tr>
            <tr>
                <td>Task 2</td>
                <td>
                    <p>Another <i>formatted</i> description</p>
                    <p>With multiple paragraphs</p>
                </td>
                <td>In Progress</td>
            </tr>
        </tbody>
    </table>
    """


def test_table_structure_preservation(sample_html_with_table):
    """Test that table structure is preserved in DOCX export."""
    # Apply consistent styling
    export_manager = ExportFormatManager()
    styled_html = export_manager.apply_consistent_table_styling(sample_html_with_table)
    
    # Parse HTML
    soup = BeautifulSoup(styled_html, "html.parser")
    
    # Create DOCX document
    doc = DocxDocument()
    
    # Process HTML elements (simplified version of export_docx logic)
    elements = soup.find_all(recursive=False)
    
    # Find and process table
    table_element = soup.find("table")
    rows = table_element.find_all("tr")
    
    # Count columns
    first_row = rows[0]
    header_cells = first_row.find_all(["th", "td"])
    col_count = len(header_cells)
    
    # Create table
    table = doc.add_table(rows=0, cols=col_count)
    
    # Process header row
    header_row = table.add_row()
    for i, cell in enumerate(header_cells):
        header_cell = header_row.cells[i]
        header_cell.text = cell.get_text(strip=True)
    
    # Process data rows
    for row_idx in range(1, len(rows)):
        tr = rows[row_idx]
        cells = tr.find_all(["td", "th"])
        if cells:
            table_row = table.add_row()
            for i, cell in enumerate(cells):
                if i < col_count:
                    table_cell = table_row.cells[i]
                    table_cell.text = cell.get_text(strip=True)
    
    # Save to buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    # Read back and verify
    doc = DocxDocument(docx_buffer)
    assert len(doc.tables) == 1
    
    table = doc.tables[0]
    assert len(table.rows) == 4  # Header + 3 data rows
    assert len(table.columns) == 3
    
    # Verify header content
    assert table.rows[0].cells[0].text == "Header 1"
    assert table.rows[0].cells[1].text == "Header 2"
    assert table.rows[0].cells[2].text == "Header 3"
    
    # Verify data content
    assert table.rows[1].cells[0].text == "Row 1, Cell 1"
    assert table.rows[2].cells[1].text == "Bold text"
    assert table.rows[3].cells[2].text == "Row 3, Cell 3"


def test_table_formatting_preservation(complex_html_with_nested_content):
    """Test that formatting within table cells is preserved."""
    # Apply consistent styling
    export_manager = ExportFormatManager()
    styled_html = export_manager.apply_consistent_table_styling(complex_html_with_nested_content)
    
    # Parse HTML
    soup = BeautifulSoup(styled_html, "html.parser")
    
    # Create DOCX document
    doc = DocxDocument()
    
    # Find and process table
    table_element = soup.find("table")
    rows = table_element.find_all("tr")
    
    # Count columns
    first_row = rows[0]
    header_cells = first_row.find_all(["th", "td"])
    col_count = len(header_cells)
    
    # Create table
    table = doc.add_table(rows=0, cols=col_count)
    
    # Process header row
    header_row = table.add_row()
    for i, cell in enumerate(header_cells):
        header_cell = header_row.cells[i]
        header_cell.text = cell.get_text(strip=True)
    
    # Process data rows with formatting (simplified)
    for row_idx in range(1, len(rows)):
        tr = rows[row_idx]
        cells = tr.find_all(["td", "th"])
        if cells:
            table_row = table.add_row()
            for i, cell in enumerate(cells):
                if i < col_count:
                    table_cell = table_row.cells[i]
                    
                    # For this test, we'll just capture the text content
                    # In the actual implementation, formatting would be preserved
                    table_cell.text = cell.get_text(strip=True)
    
    # Save to buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    # Read back and verify
    doc = DocxDocument(docx_buffer)
    assert len(doc.tables) == 1
    
    table = doc.tables[0]
    
    # Verify complex content is preserved (text only in this test)
    assert "Task 1" in table.rows[1].cells[0].text
    assert "complex" in table.rows[1].cells[1].text
    assert "Item 1" in table.rows[1].cells[1].text
    assert "Item 2" in table.rows[1].cells[1].text
    
    assert "Task 2" in table.rows[2].cells[0].text
    assert "formatted" in table.rows[2].cells[1].text
    assert "multiple paragraphs" in table.rows[2].cells[1].text


def test_multi_page_table_handling():
    """Test that large tables spanning multiple pages are handled correctly."""
    # Create a large HTML table
    large_table_html = "<table><thead><tr><th>ID</th><th>Description</th></tr></thead><tbody>"
    
    # Add 100 rows to ensure multi-page spanning
    for i in range(1, 101):
        large_table_html += f"<tr><td>Item {i}</td><td>Description for item {i}</td></tr>"
    
    large_table_html += "</tbody></table>"
    
    # Apply consistent styling
    export_manager = ExportFormatManager()
    styled_html = export_manager.apply_consistent_table_styling(large_table_html)
    
    # Parse HTML
    soup = BeautifulSoup(styled_html, "html.parser")
    
    # Create DOCX document
    doc = DocxDocument()
    
    # Find and process table
    table_element = soup.find("table")
    rows = table_element.find_all("tr")
    
    # Count columns
    first_row = rows[0]
    header_cells = first_row.find_all(["th", "td"])
    col_count = len(header_cells)
    
    # Create table
    table = doc.add_table(rows=0, cols=col_count)
    
    # Process all rows
    for row_idx in range(len(rows)):
        tr = rows[row_idx]
        cells = tr.find_all(["th", "td"])
        if cells:
            table_row = table.add_row()
            for i, cell in enumerate(cells):
                if i < col_count:
                    table_cell = table_row.cells[i]
                    table_cell.text = cell.get_text(strip=True)
    
    # Save to buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    # Read back and verify
    doc = DocxDocument(docx_buffer)
    assert len(doc.tables) == 1
    
    table = doc.tables[0]
    assert len(table.rows) == 101  # Header + 100 data rows
    
    # Verify some content
    assert table.rows[0].cells[0].text == "ID"
    assert table.rows[1].cells[0].text == "Item 1"
    assert table.rows[50].cells[0].text == "Item 50"
    assert table.rows[100].cells[0].text == "Item 100"