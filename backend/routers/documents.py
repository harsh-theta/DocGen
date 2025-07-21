from fastapi import APIRouter, UploadFile, File, HTTPException, Request
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime
from backend.core.config import settings
import re
from storage3.exceptions import StorageApiError
import uuid
from backend.models.document import Document
from backend.core.database import get_db
from sqlalchemy.ext.asyncio import AsyncSession
from fastapi import Depends
from backend.services.auth import get_current_user
from backend.models.user import User
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
import io
from docx import Document as DocxDocument
from docx.shared import Inches
import html2text
import tempfile
import os
import logging
import time
import random
from backend.ai.context_handler import ProjectContextHandler
from backend.ai.workflow import DocumentGenerationWorkflow
from backend.ai.models import GenerationRequest, GenerationResponse, create_generation_result
import asyncio
import json
import pdfplumber
from pathlib import Path
from docx.table import _Cell
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from enum import Enum
from bs4 import BeautifulSoup

# Initialize logger for this module
logger = logging.getLogger(__name__)

router = APIRouter()

dummy_documents = [
    {
        "id": "1",
        "name": "Sample.docx",
        "created_at": datetime.utcnow().isoformat(),
        "type": "original",
        "status": "Uploaded",
        "original_file_url": "https://example.com/sample.docx",
        "ai_content": None,
        "final_file_url": None,
        "parsed_structure": "<p>Sample content</p>",
    },
    {
        "id": "2",
        "name": "Generated.pdf",
        "created_at": datetime.utcnow().isoformat(),
        "type": "generated",
        "status": "Generated",
        "original_file_url": None,
        "ai_content": "<p>Generated content</p>",
        "final_file_url": "https://example.com/generated.pdf",
        "parsed_structure": None,
    },
]

class DocumentResponse(BaseModel):
    id: str
    name: str
    created_at: str
    type: str
    status: Optional[str]
    original_file_url: Optional[str]
    ai_content: Optional[str]
    final_file_url: Optional[str]
    parsed_structure: Optional[str]

# Add a placeholder parsing utility
from backend.ai.html_parser import HTMLSectionParser
import os

def para_to_html(para):
    text = para.text.strip()
    if not text:
        return ""
    style = para.style.name.lower() if para.style else ""
    # Inline formatting
    html = ""
    for run in para.runs:
        run_text = run.text.replace("<", "&lt;").replace(">", "&gt;")
        if not run_text:
            continue
        if run.bold:
            run_text = f"<b>{run_text}</b>"
        if run.italic:
            run_text = f"<i>{run_text}</i>"
        if run.underline:
            run_text = f"<u>{run_text}</u>"
        html += run_text
    if style.startswith("heading"):
        level = ''.join(filter(str.isdigit, style)) or "2"
        return f"<h{level}>{html}</h{level}>"
    elif style.startswith("list") or para._element.xpath('.//w:numPr'):
        # Will be handled in list logic
        return None
    else:
        return f"<p>{html}</p>"

def table_to_html(table):
    html = "<table>"
    for row in table.rows:
        html += "<tr>"
        for cell in row.cells:
            html += f"<td>{cell.text.strip()}</td>"
        html += "</tr>"
    html += "</table>"
    return html

def docx_to_html(doc):
    html_parts = []
    block_items = []
    # Collect paragraphs and tables in order
    for block in doc.element.body:
        if isinstance(block, CT_P):
            block_items.append(Paragraph(block, doc))
        elif isinstance(block, CT_Tbl):
            block_items.append(doc.tables[len([b for b in block_items if isinstance(b, _Cell)])])
    i = 0
    while i < len(block_items):
        item = block_items[i]
        if isinstance(item, Paragraph):
            # Handle lists
            if item._element.xpath('.//w:numPr'):
                # Start of a list
                list_tag = "ul" if item.style.name.lower().find("number") == -1 else "ol"
                html_parts.append(f"<{list_tag}>")
                while i < len(block_items) and isinstance(block_items[i], Paragraph) and block_items[i]._element.xpath('.//w:numPr'):
                    html_parts.append(f"<li>{block_items[i].text.strip()}</li>")
                    i += 1
                html_parts.append(f"</{list_tag}>")
                continue
            # Normal paragraph/heading
            html = para_to_html(item)
            if html:
                html_parts.append(html)
        elif hasattr(item, 'rows'):
            html_parts.append(table_to_html(item))
        i += 1
    return "\n".join(html_parts)

def parse_file_to_html(file_path: str) -> str:
    """
    Advanced parsing for DOCX and PDF:
    - DOCX: Extract headings, paragraphs, tables, lists, and inline formatting.
    - PDF: Group lines into paragraphs, try to detect headings by font size.
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".docx":
        try:
            doc = DocxDocument(file_path)
            return docx_to_html(doc)
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            return ""
    elif ext == ".pdf":
        try:
            html_parts = []
            with pdfplumber.open(file_path) as pdf:
                prev_fontsize = None
                para_lines = []
                for page in pdf.pages:
                    words = page.extract_words(extra_attrs=["size"]) if hasattr(page, 'extract_words') else []
                    lines = page.extract_text().splitlines() if page.extract_text() else []
                    # Group lines into paragraphs by empty lines
                    for line in lines:
                        if not line.strip():
                            if para_lines:
                                html_parts.append(f"<p>{' '.join(para_lines)}</p>")
                                para_lines = []
                        else:
                            para_lines.append(line.strip())
                    if para_lines:
                        html_parts.append(f"<p>{' '.join(para_lines)}</p>")
                        para_lines = []
                    # Try to detect headings by font size (if available)
                    for word in words:
                        if word.get("size", 0) > 15:  # Arbitrary threshold
                            html_parts.append(f"<h2>{word['text']}</h2>")
            return "\n".join(html_parts)
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            return ""
    else:
        logger.warning(f"Unsupported file type for parsing: {ext}")
        return ""

@router.post("/upload-doc")
async def upload_doc(request: Request, file: UploadFile = File(...), title: str = None, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required.")
    # File type validation
    allowed_extensions = {"pdf", "docx", "doc"}
    if "." not in file.filename or file.filename.rsplit(".", 1)[1].lower() not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Only .pdf, .docx, and .doc files are allowed.")
    supabase = request.app.state.supabase
    bucket = settings.SUPABASE_BUCKET
    contents = await file.read()
    # Sanitize filename and append uuid
    base, ext = (str(file.filename).rsplit('.', 1) + [""])[:2]
    safe_base = re.sub(r'[^A-Za-z0-9._-]', '_', base)
    unique_filename = f"{safe_base}_{uuid.uuid4().hex}"
    if ext:
        unique_filename += f".{ext}"
    try:
        res = supabase.storage.from_(bucket).upload(unique_filename, contents)
    except StorageApiError as e:
        error_info = e.args[0] if e.args else {}
        message = error_info.get("message") if isinstance(error_info, dict) else str(error_info)
        raise HTTPException(status_code=500, detail=f"Failed to upload file: {message}")
    public_url = supabase.storage.from_(bucket).get_public_url(unique_filename)
    user_id = current_user.id
    # Save file temporarily for parsing
    temp_path = f"/tmp/{unique_filename}"
    with open(temp_path, "wb") as f:
        f.write(contents)
    # Parse the file to HTML
    try:
        parsed_html = parse_file_to_html(temp_path)
    except Exception as e:
        parsed_html = None
        logger.error(f"Failed to parse uploaded file to HTML: {e}")
    # Clean up temp file
    try:
        os.remove(temp_path)
    except Exception:
        pass
    # Use project name as title if provided
    doc_title = title or file.filename
    doc = Document(
        user_id=user_id,
        name=file.filename,
        stored_filename=unique_filename,
        created_at=datetime.utcnow(),
        type="original",
        status="Uploaded",
        original_file_url=public_url,
        ai_content=None,
        final_file_url=None,
        parsed_structure=parsed_html,
        # Add title field if present
        title=doc_title if hasattr(Document, 'title') else None
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return {
        "id": str(doc.id),
        "name": doc.name,
        "title": getattr(doc, 'title', None) or doc.name,
        "created_at": doc.created_at.isoformat(),
        "type": doc.type,
        "status": doc.status,
        "original_file_url": doc.original_file_url,
        "ai_content": doc.ai_content,
        "final_file_url": doc.final_file_url,
        "parsed_structure": doc.parsed_structure,
        "stored_filename": doc.stored_filename
    }

@router.get("/documents", response_model=List[DocumentResponse])
async def list_documents(db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    user_id = current_user.id
    result = await db.execute(
        Document.__table__.select().where(Document.user_id == user_id)
    )
    docs = result.fetchall()
    return [
        {
            "id": str(row.id),
            "name": row.name,
            "created_at": row.created_at.isoformat() if row.created_at else None,
            "type": row.type,
            "status": row.status,
            "original_file_url": row.original_file_url,
            "ai_content": row.ai_content,
            "final_file_url": row.final_file_url,
            "parsed_structure": row.parsed_structure,
            "stored_filename": row.stored_filename
        }
        for row in docs
    ]

@router.get("/session/{id}")
async def get_document(id: str, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    user_id = current_user.id
    result = await db.execute(
        Document.__table__.select().where(Document.id == id).where(Document.user_id == user_id)
    )
    row = result.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    return {
        "id": str(row.id),
        "name": row.name,
        "created_at": row.created_at.isoformat() if row.created_at else None,
        "type": row.type,
        "status": row.status,
        "original_file_url": row.original_file_url,
        "ai_content": row.ai_content,
        "final_file_url": row.final_file_url,
        "parsed_structure": row.parsed_structure,
        "stored_filename": row.stored_filename
    }

class SaveEditsRequest(BaseModel):
    id: str
    content: str

@router.post("/save-edits")
async def save_edits(data: SaveEditsRequest, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    user_id = current_user.id
    result = await db.execute(
        Document.__table__.select().where(Document.id == data.id).where(Document.user_id == user_id)
    )
    row = result.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    await db.execute(
        Document.__table__.update().where(Document.id == data.id).where(Document.user_id == user_id).values(ai_content=data.content)
    )
    await db.commit()
    return {"message": "Edits saved"}

# New endpoint to directly download PDF files
@router.get("/download/pdf/{id}")
async def download_pdf(id: str, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    """
    Direct PDF download endpoint that serves the file with proper headers
    to force download instead of viewing in browser
    """
    request_id = f"pdf_download_{id}_{uuid.uuid4().hex[:8]}"
    logger.info(f"[{request_id}] Starting direct PDF download for document {id}")
    try:
        from backend.services.pdf_generator import (
            generate_pdf_from_html,
            PDFGenerationError
        )
        user_id = current_user.id
        result = await db.execute(
            Document.__table__.select().where(Document.id == id).where(Document.user_id == user_id)
        )
        row = result.fetchone()
        if not row:
            logger.warning(f"[{request_id}] Document not found for user {user_id}, document {id}")
            raise HTTPException(status_code=404, detail="Document not found")
        content = row.ai_content
        if not content or not content.strip():
            content = "<p>No content available for this document.</p>"
        document_title = row.name or "Document"
        logger.debug(f"[{request_id}] Generating PDF for direct download")
        pdf_bytes = await generate_pdf_from_html(content, document_title, document_id=id, user_id=user_id)
        import tempfile
        import os
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        temp_file.write(pdf_bytes)
        temp_file.close()
        safe_filename = re.sub(r'[^A-Za-z0-9._-]', '_', document_title)
        download_filename = f"{safe_filename}.pdf"
        logger.info(f"[{request_id}] Serving PDF file directly: {download_filename}")
        # Check if file exists before returning
        if not os.path.exists(temp_file.name):
            logger.error(f"[{request_id}] PDF file not found on disk: {temp_file.name}")
            raise HTTPException(status_code=404, detail="PDF file not found. Please try exporting again.")
        return FileResponse(
            path=temp_file.name,
            filename=download_filename,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{download_filename}"',
                "Content-Type": "application/pdf"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[{request_id}] Error serving PDF file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")

@router.post("/export/pdf/{id}")
async def export_pdf(id: str, request: Request, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        user_id = current_user.id
        result = await db.execute(
            Document.__table__.select().where(Document.id == id).where(Document.user_id == user_id)
        )
        row = result.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Document not found")
        content = row.ai_content or "<p>No content available</p>"
        document_title = getattr(row, 'title', None) or row.name or 'Document'
        logger.info(f"[PDF Export] Generating PDF for document {id} with title '{document_title}'")
        # Inject default font into HTML
        html_with_font = f"<style>body, p, h1, h2, h3, h4, h5, h6, td, th {{ font-family: Arial, sans-serif !important; }}</style>" + content
        from backend.services.pdf_generator import generate_pdf_from_html
        try:
            pdf_bytes = await generate_pdf_from_html(html_with_font, document_title, document_id=id, user_id=user_id)
        except Exception as e:
            logger.error(f"[PDF Export] PDF generation failed: {e}")
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")
        import tempfile
        import os
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        temp_file.write(pdf_bytes)
        temp_file.close()
        safe_filename = re.sub(r'[^A-Za-z0-9._-]', '_', document_title)
        download_filename = f"{safe_filename}_{id}_{uuid.uuid4().hex}.pdf"
        supabase = request.app.state.supabase
        bucket = settings.SUPABASE_BUCKET
        try:
            with open(temp_file.name, "rb") as f:
                supabase.storage.from_(bucket).upload(download_filename, f.read())
            public_url = supabase.storage.from_(bucket).get_public_url(download_filename)
        except Exception as e:
            logger.error(f"[PDF Export] PDF upload failed: {e}")
            raise HTTPException(status_code=500, detail=f"PDF upload failed: {e}")
        await db.execute(
            Document.__table__.update().where(Document.id == id).where(Document.user_id == user_id).values(final_file_url=public_url)
        )
        await db.commit()
        # Always return the Supabase public URL for download
        return {"url": public_url}
    except Exception as e:
        logger.error(f"[PDF Export] Error: {e}")
        raise HTTPException(status_code=500, detail=f"PDF export failed: {e}")

@router.post("/export/docx/{id}")
async def export_docx(id: str, request: Request, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        user_id = current_user.id
        result = await db.execute(
            Document.__table__.select().where(Document.id == id).where(Document.user_id == user_id)
        )
        row = result.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Document not found")
        content = row.ai_content or "<p>No content available</p>"
        document_title = getattr(row, 'title', None) or row.name or 'Document'
        h = BeautifulSoup(content, "html.parser")
        docx_buffer = io.BytesIO()
        doc = DocxDocument()
        # Set default font to Arial
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = None  # Use default size
        doc.add_heading(document_title, 0)
        for el in h.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol", "li"]):
            if el.name == "h1":
                doc.add_heading(el.get_text(strip=True), level=1)
            elif el.name == "h2":
                doc.add_heading(el.get_text(strip=True), level=2)
            elif el.name == "h3":
                doc.add_heading(el.get_text(strip=True), level=3)
            elif el.name == "h4":
                doc.add_heading(el.get_text(strip=True), level=4)
            elif el.name == "h5":
                doc.add_heading(el.get_text(strip=True), level=5)
            elif el.name == "h6":
                doc.add_heading(el.get_text(strip=True), level=6)
            elif el.name == "p":
                doc.add_paragraph(el.get_text(strip=True))
            elif el.name in ["ul", "ol"]:
                for li in el.find_all("li"):
                    doc.add_paragraph(li.get_text(strip=True), style="List Bullet" if el.name == "ul" else "List Number")
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        supabase = request.app.state.supabase
        bucket = settings.SUPABASE_BUCKET
        safe_filename = re.sub(r'[^A-Za-z0-9._-]', '_', document_title)
        filename = f"{safe_filename}_{id}_{uuid.uuid4().hex}.docx"
        try:
            supabase.storage.from_(bucket).upload(filename, docx_buffer.read())
            public_url = supabase.storage.from_(bucket).get_public_url(filename)
        except Exception as e:
            logger.error(f"[DOCX Export] DOCX upload failed: {e}")
            raise HTTPException(status_code=500, detail=f"DOCX upload failed: {e}")
        await db.execute(
            Document.__table__.update().where(Document.id == id).where(Document.user_id == user_id).values(final_file_url=public_url)
        )
        await db.commit()
        return {"url": public_url}
    except Exception as e:
        logger.error(f"[DOCX Export] Error: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {e}")

@router.post("/export/markdown/{id}")
async def export_markdown(id: str, request: Request, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    user_id = current_user.id
    result = await db.execute(
        Document.__table__.select().where(Document.id == id).where(Document.user_id == user_id)
    )
    row = result.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    content = row.ai_content or ""
    # Save as Markdown file (assume content is already Markdown or plain text)
    md_bytes = content.encode("utf-8")
    supabase = request.app.state.supabase
    bucket = settings.SUPABASE_BUCKET
    filename = f"exported_{id}_{uuid.uuid4().hex}.md"
    supabase.storage.from_(bucket).upload(filename, md_bytes)
    public_url = supabase.storage.from_(bucket).get_public_url(filename)
    # Update DB
    await db.execute(
        Document.__table__.update().where(Document.id == id).where(Document.user_id == user_id).values(final_file_url=public_url)
    )
    await db.commit()
    return {"url": public_url} 

# Helper to make GeneratedSection JSON serializable

def section_to_dict(section):
    d = section.__dict__.copy()
    for k, v in d.items():
        if isinstance(v, Enum):
            d[k] = v.value
    return d

@router.post("/generate", response_model=GenerationResponse)
async def generate_document(
    request: Request,
    data: GenerationRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    AI-powered document generation endpoint.
    Accepts HTML template and user input, runs the AI workflow, saves result.
    """
    logger.info(f"[AI] Starting document generation for user {current_user.id}, doc {data.document_id}")
    try:
        # Validate and build context
        context = ProjectContextHandler.build_project_context(data.user_input)
        err = ProjectContextHandler.validate_project_context(context)
        if err:
            logger.warning(f"[AI] Invalid input: {err}")
            return GenerationResponse(success=False, errors=[err])
        # Run workflow
        workflow = DocumentGenerationWorkflow()
        state = await workflow.run(data.html_template, context)
        # Prepare fields to save
        result_html = state.get("final_html", "")
        if result_html is None:
            result_html = ""
        errors = state.get("errors", [])
        # New AI fields
        html_template = data.html_template
        ai_generation_metadata = state.get("metadata", {})
        generated_sections = state.get("generated_sections", [])
        # Convert to JSON strings if needed
        ai_generation_metadata_str = json.dumps(ai_generation_metadata) if not isinstance(ai_generation_metadata, str) else ai_generation_metadata
        # Use section_to_dict for JSON serialization
        generated_sections_str = json.dumps([section_to_dict(g) for g in generated_sections]) if generated_sections and not isinstance(generated_sections, str) else (generated_sections if isinstance(generated_sections, str) else "[]")
        # Update document record
        user_id = current_user.id
        result = await db.execute(
            Document.__table__.select().where(Document.id == data.document_id).where(Document.user_id == user_id)
        )
        row = result.fetchone()
        if not row:
            logger.warning(f"[AI] Document not found for user {user_id}, doc {data.document_id}")
            return GenerationResponse(success=False, errors=["Document not found"]) 
        await db.execute(
            Document.__table__.update().where(Document.id == data.document_id).where(Document.user_id == user_id).values(
                ai_content=result_html,
                html_template=html_template,
                ai_generation_metadata=ai_generation_metadata_str,
                generated_sections=generated_sections_str
            )
        )
        await db.commit()
        logger.info(f"[AI] Document generation complete for doc {data.document_id}")
        return GenerationResponse(
            success=True,
            generated_html=result_html,
            sections_processed=state.get("total_sections", 0),
            sections_failed=len(errors),
            errors=errors,
            generation_time_ms=state.get("metadata", {}).get("generation_time_ms", 0),
            status="completed" if not errors else "partial",
            metadata=state.get("metadata", {})
        )
    except Exception as e:
        logger.error(f"[AI] Error in document generation: {str(e)}", exc_info=True)
        return GenerationResponse(success=False, errors=[str(e)]) 